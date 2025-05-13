from flask import Flask, request, send_file, render_template_string
from pdf2docx import Converter
from docx import Document
from PIL import Image, ImageDraw
from pdf2image import convert_from_path
import easyocr
import os
import uuid
import zipfile

app = Flask(__name__)

HTML_PAGE = """
<!DOCTYPE html>
<html>
<head>
  <title>CloudConverterX</title>
  <style>
    body {
      font-family: 'Segoe UI', sans-serif;
      background-color: #f0f8ff;
      text-align: center;
      padding-top: 40px;
    }
    form {
      background: #ffffff;
      padding: 30px;
      display: inline-block;
      border-radius: 12px;
      box-shadow: 0 0 12px rgba(0,0,0,0.2);
    }
    h1 {
      color: #333366;
    }
    select, input[type="file"], button {
      font-size: 16px;
      padding: 10px;
      margin: 10px;
    }
    button {
      background-color: #4CAF50;
      color: white;
      border: none;
      border-radius: 6px;
    }
  </style>
</head>
<body>
  <h1>Cloud File Converter</h1>
  <form action="/convert" method="POST" enctype="multipart/form-data">
    <label><b>Select File:</b></label><br>
    <input type="file" name="file" required><br><br>

    <label><b>Conversion Type:</b></label><br>
    <select name="option" required>
      <option value="pdf2docx">PDF to Word</option>
      <option value="docx2pdf">Word to PDF</option>
      <option value="image2word">Image to Word (OCR)</option>
      <option value="word2image">Word to Image</option>
      <option value="image2pdf">Image to PDF</option>
      <option value="pdf2image">PDF to Image</option>
    </select><br><br>

    <button type="submit">Convert</button>
  </form>
</body>
</html>
"""

@app.route('/')
def index():
    return render_template_string(HTML_PAGE)

@app.route('/convert', methods=['POST'])
def convert():
    file = request.files['file']
    option = request.form['option']
    filename = str(uuid.uuid4()) + "_" + file.filename
    file.save(filename)

    try:
        if option == 'pdf2docx':
            output = filename.replace(".pdf", ".docx")
            cv = Converter(filename)
            cv.convert(output)
            cv.close()

        elif option == 'docx2pdf':
            try:
                from docx2pdf import convert
                output = filename.replace(".docx", ".pdf")
                convert(filename, output)
            except:
                # Fallback: create a simple PDF with text
                from docx import Document
                from reportlab.pdfgen import canvas
                temp_doc = Document(filename)
                text = "\n".join([p.text for p in temp_doc.paragraphs])
                output = filename.replace(".docx", ".pdf")
                c = canvas.Canvas(output)
                for i, line in enumerate(text.splitlines()):
                    c.drawString(50, 800 - 15*i, line)
                c.save()

        elif option == 'image2word':
            reader = easyocr.Reader(['en'])
            result = reader.readtext(filename, detail=0)
            doc = Document()
            doc.add_heading('Extracted Text', 0)
            for line in result:
                doc.add_paragraph(line)
            output = filename.rsplit('.', 1)[0] + "_ocr.docx"
        
        elif option == 'word2image':
            doc = Document(filename)
            text = "\n".join([p.text for p in doc.paragraphs])
            img = Image.new("RGB", (800, 1000), color="white")
            draw = ImageDraw.Draw(img)
            draw.text((20, 20), text[:4000], fill="black")
            output = filename.replace(".docx", ".png")
            img.save(output)

        elif option == 'image2pdf':
            img = Image.open(filename).convert('RGB')
            output = filename.rsplit('.', 1)[0] + ".pdf"
            img.save(output)

        elif option == 'pdf2image':
            images = convert_from_path(filename)
            output_files = []
            for i, img in enumerate(images):
                out = f"{filename}_page{i+1}.png"
                img.save(out, 'PNG')
                output_files.append(out)

            zipname = filename.replace(".pdf", "_images.zip")
            with zipfile.ZipFile(zipname, 'w') as zipf:
                for f in output_files:
                    zipf.write(f)
                    os.remove(f)
            output = zipname

        else:
            return "Unsupported option", 400

        os.remove(filename)
        return send_file(output, as_attachment=True)

    except Exception as e:
        os.remove(filename)
        return f"Error: {str(e)}", 500

if __name__ == '__main__':
    app.run(debug=True)
